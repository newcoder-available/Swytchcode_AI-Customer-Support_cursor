"""Generate ResolveAI project submission Word documentation."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=18 if level == 1 else 14 if level == 2 else 12, bold=True)
    return h


def add_para(doc, text, bold=False, size=11, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ResolveAI")
    set_run_font(run, size=28, bold=True, color=RGBColor(15, 118, 110))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "AI Customer Support Agent Powered by Swytchcode & Gmail"
    )
    set_run_font(run, size=14, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Project Submission Documentation\n"
        "End-to-End Features, Architecture, Coding & Usage Guide\n"
        "Repository: https://github.com/newcoder-available/Swytchcode_AI-Customer-Support_cursor"
    )
    set_run_font(run, size=11)

    add_para(doc, "")

    # 1. Purpose
    add_heading(doc, "1. Project Purpose and Goals", 1)
    add_para(
        doc,
        "ResolveAI is an AI-powered customer support workspace that turns Gmail "
        "conversations into live support tickets and automates the resolution workflow "
        "through Swytchcode. The application does not call Gmail APIs directly from the "
        "frontend. All external actions go through Swytchcode as the execution kernel.",
    )
    add_para(doc, "Goals:", bold=True)
    add_bullets(
        doc,
        [
            "Answer support emails using grounded AI responses.",
            "Search company knowledge before replying (no hallucinated procedures).",
            "Draft and send customer updates automatically when confidence is high.",
            "Create support tickets from user input with customer email notification.",
            "Escalate low-confidence or high-risk issues instead of guessing.",
            "Provide a production-style Inbox / Tickets UI for operators.",
        ],
    )

    # 2. Problem statement
    add_heading(doc, "2. Problem Statement Alignment", 1)
    add_para(
        doc,
        "Hackathon brief: Answer support emails, search company knowledge, draft replies, "
        "create engineering tickets, and send customer updates. Integrations listed in the "
        "brief include Gmail, Notion, Jira, Resend, and GitHub.",
    )
    add_para(doc, "Current MVP alignment:", bold=True)
    add_table(
        doc,
        ["Capability", "Status", "Implementation"],
        [
            [
                "Answer support emails",
                "Implemented (Gmail)",
                "Inbox + thread reply via gmail.user.send.create1",
            ],
            [
                "Search company knowledge",
                "Implemented (local KB)",
                "Markdown knowledge base + retrieve/ground pipeline",
            ],
            [
                "Draft replies",
                "Implemented",
                "runSupportAgent: classify → retrieve → ground → confidence",
            ],
            [
                "Create tickets",
                "Implemented (Gmail tickets)",
                "Create form + Gmail thread ID as ticket ID",
            ],
            [
                "Send customer updates",
                "Implemented (Gmail)",
                "Auto notify on create + AI update on ticket detail",
            ],
            [
                "Notion / Jira / Resend / GitHub",
                "Not yet wired",
                "Available in Swytchcode registry for future expansion",
            ],
        ],
    )

    # 3. Features
    add_heading(doc, "3. Key Features", 1)
    add_heading(doc, "3.1 Gmail Support Inbox", 2)
    add_bullets(
        doc,
        [
            "Live Gmail OAuth via Swytchcode (swytchcode auth connect Gmail).",
            "Gmail thread = support ticket (threadId is the ticket ID).",
            "List ResolveAI-related threads and merge app-created tickets into Inbox.",
            "Open conversation, distinguish customer vs ResolveAI vs human messages.",
            "Polling refresh (configurable GMAIL_POLL_INTERVAL_MS).",
        ],
    )
    add_heading(doc, "3.2 Create Ticket from User Input", 2)
    add_bullets(
        doc,
        [
            "Form fields: customer name, customer email (required), subject, issue, priority.",
            "Creates a Gmail thread and emails the customer a ticket-created notification.",
            "Customer email is stored and shown on the ticket list and detail pages.",
            "Runs the AI agent immediately after creation.",
            "Sends an automatic AI follow-up/update on the same Gmail thread.",
        ],
    )
    add_heading(doc, "3.3 AI Agent Pipeline", 2)
    add_bullets(
        doc,
        [
            "Intent classification (troubleshooting, knowledge, escalation, ticket create/status).",
            "Knowledge retrieval from local company KB articles.",
            "Grounded answer generation with cited sources.",
            "Confidence scoring with policy thresholds (auto-reply / clarify / escalate).",
            "Idempotent processing to avoid duplicate replies to the same message.",
        ],
    )
    add_heading(doc, "3.4 Ticket Detail & AI Updates", 2)
    add_bullets(
        doc,
        [
            "Click any ticket row to open the detail page.",
            "View customer email, status, activity, knowledge sources, latest AI answer.",
            "Send AI update to customer (regenerates grounded response and emails it).",
            "Send manual update text to the same Gmail thread.",
            "Open ticket in Inbox via deep link (?thread=id).",
        ],
    )
    add_heading(doc, "3.5 Status Engine", 2)
    add_bullets(
        doc,
        [
            "NEW — new customer/support item with no AI response yet.",
            "AI_ANALYZING — agent processing.",
            "AI_RESPONDED — ResolveAI replied; waiting on customer.",
            "WAITING_FOR_CUSTOMER — clarification requested.",
            "CUSTOMER_REPLIED — new customer message after AI response.",
            "ESCALATED — low confidence or critical path.",
            "RESOLVED — marked resolved.",
        ],
    )

    # 4. Architecture
    add_heading(doc, "4. System Architecture", 1)
    add_para(doc, "High-level flow:", bold=True)
    add_para(
        doc,
        "Operator / Customer Email → ResolveAI UI (Next.js) → API Routes → "
        "Gmail Service / Agent → Swytchcode Runtime (@swytchcode/runtime) → "
        "tooling.json allowlist → Swytchcode Kernel → Gmail → Structured Result → UI",
    )
    add_para(doc, "Ticket create architecture:", bold=True)
    add_para(
        doc,
        "User input → createSupportTicketFromInput() → Gmail send (notification) → "
        "runSupportAgent() → Gmail send (AI update on same thread) → store ticket → "
        "show in Tickets + Inbox",
    )
    add_para(doc, "Security principles:", bold=True)
    add_bullets(
        doc,
        [
            "No Gmail passwords, OAuth tokens, or API keys in the frontend.",
            "Credentials managed by Swytchcode auth; live mode must not override OAuth with fake Bearer tokens.",
            "Only allowlisted Swytchcode operations can run.",
            ".env.local is never committed.",
            "No delete/trash/arbitrary mailbox mutation operations enabled.",
        ],
    )

    # 5. Tech stack
    add_heading(doc, "5. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Frontend / App", "Next.js 15 (App Router), React 19, TypeScript"],
            ["Styling", "Tailwind CSS + custom product CSS"],
            ["Integrations", "Swytchcode CLI + @swytchcode/runtime"],
            ["Email / Tickets", "Gmail via Swytchcode OAuth"],
            ["Knowledge", "Local markdown KB under /knowledge"],
            ["Persistence", "Local JSON ticket store (data/gmail-ticket-store.json)"],
            ["Validation", "Zod + custom validation helpers"],
        ],
    )

    # 6. Code structure
    add_heading(doc, "6. Code Structure", 1)
    add_table(
        doc,
        ["Path", "Responsibility"],
        [
            ["src/app/(product)/inbox", "Support inbox UI page"],
            ["src/app/(product)/tickets", "Ticket list + create form + detail"],
            ["src/app/api/gmail/*", "Inbox, poll, thread, reply, process APIs"],
            ["src/app/api/tickets/*", "Create/list tickets + AI update API"],
            ["src/components/inbox", "3-pane InboxWorkspace"],
            ["src/components/tickets", "TicketsWorkspace create/list UI"],
            ["src/lib/agent/*", "classify, retrieve, ground, confidence, answer"],
            ["src/lib/gmail/*", "client, service, create, process, store, status"],
            ["src/lib/swytchcode/*", "allowlist, exec, health, actions"],
            ["knowledge/*", "Company knowledge articles"],
            ["docs/GMAIL_SWYTCHCODE.md", "Canonical ID mapping"],
            [".swytchcode/tooling.json", "Enabled Swytchcode methods"],
        ],
    )

    # 7. Swytchcode ops
    add_heading(doc, "7. Swytchcode Gmail Operations", 1)
    add_para(
        doc,
        "Operations were discovered with `swytchcode get gmail` and `swytchcode info` "
        "(not invented). Enabled methods:",
    )
    add_table(
        doc,
        ["App Operation", "Canonical ID", "Purpose"],
        [
            ["gmail.profile", "gmail.user.profile.get", "Mailbox profile / connection"],
            ["gmail.threads.list", "gmail.user.threads.get", "List/search support threads"],
            ["gmail.threads.get", "gmail.user.threads.get1", "Read full ticket thread"],
            ["gmail.messages.get", "gmail.user.messages.get1", "Read one message"],
            ["gmail.messages.list", "gmail.user.messages.get", "List messages"],
            ["gmail.messages.send", "gmail.user.send.create1", "Send/reply (raw + threadId)"],
            ["gmail.labels.list", "gmail.user.labels.get", "List labels"],
            ["gmail.threads.modify", "gmail.user.modify.create1", "Modify thread labels"],
        ],
    )

    # 8. Setup
    add_heading(doc, "8. Setup and Configuration", 1)
    add_heading(doc, "8.1 Prerequisites", 2)
    add_bullets(
        doc,
        [
            "Node.js 20+",
            "Global Swytchcode CLI (`npm install -g swytchcode`)",
            "Swytchcode login (`swytchcode login`)",
            "Gmail account for support inbox OAuth",
        ],
    )
    add_heading(doc, "8.2 Install Project", 2)
    add_numbered(
        doc,
        [
            "Clone: https://github.com/newcoder-available/Swytchcode_AI-Customer-Support_cursor",
            "cd into the project folder",
            "npm install",
            "Copy .env.example to .env.local and edit values",
        ],
    )
    add_heading(doc, "8.3 Enable Gmail in Swytchcode", 2)
    add_numbered(
        doc,
        [
            "swytchcode get gmail",
            "swytchcode add method gmail.user.profile.get",
            "swytchcode add method gmail.user.threads.get",
            "swytchcode add method gmail.user.threads.get1",
            "swytchcode add method gmail.user.messages.get",
            "swytchcode add method gmail.user.messages.get1",
            "swytchcode add method gmail.user.send.create1",
            "swytchcode add method gmail.user.labels.get",
            "swytchcode add method gmail.user.modify.create1",
            "swytchcode auth connect Gmail",
            "swytchcode list tooling  (confirm methods enabled)",
            "swytchcode auth status   (Gmail should show connected)",
        ],
    )
    add_heading(doc, "8.4 Environment Variables (.env.local)", 2)
    add_table(
        doc,
        ["Variable", "Example / Notes"],
        [
            ["SWYTCHCODE_MODE", "live (required for real Gmail)"],
            ["SUPPORT_INBOX_EMAIL", "your authenticated Gmail address"],
            ["SUPPORT_LABEL", "ResolveAI (optional filter)"],
            ["SUPPORT_LABEL_REQUIRED", "false (recommended for MVP)"],
            ["GMAIL_USER_ID", "me"],
            ["GMAIL_POLL_INTERVAL_MS", "15000"],
            ["GMAIL_AUTO_REPLY_CONFIDENCE", "0.85"],
            ["GMAIL_CLARIFY_CONFIDENCE", "0.6"],
        ],
    )
    add_heading(doc, "8.5 Run", 2)
    add_numbered(
        doc,
        [
            "npm run dev",
            "Open http://localhost:3000/inbox",
            "Open http://localhost:3000/tickets for create/update workflow",
        ],
    )

    # 9. End-to-end usage
    add_heading(doc, "9. End-to-End Usage Guide", 1)
    add_heading(doc, "9.1 Create Ticket + Auto Notify + Auto AI Update", 2)
    add_numbered(
        doc,
        [
            "Go to Tickets page.",
            "Enter Customer Name, Customer Email, Subject, Issue, Priority.",
            "Click “Create ticket & notify email”.",
            "System creates a Gmail thread and emails the customer a creation notification.",
            "AI agent analyzes the issue against the knowledge base.",
            "System automatically emails an AI update/reply on the same thread.",
            "Ticket appears in Tickets list with customer email visible and Notify=Sent.",
            "Ticket also appears in Inbox.",
        ],
    )
    add_heading(doc, "9.2 Open Existing Ticket and Send Another AI Update", 2)
    add_numbered(
        doc,
        [
            "On Tickets page, click the ticket row (or ticket ID link).",
            "Review customer email, status, activity, and previous AI answer.",
            "Optionally type a note for the AI.",
            "Click “Send AI update to customer”.",
            "Customer receives an update email on the same Gmail thread.",
            "Or click “Send manual update” to send your own text.",
        ],
    )
    add_heading(doc, "9.3 Inbox Conversation Workflow", 2)
    add_numbered(
        doc,
        [
            "Open Inbox.",
            "Select a ticket from the left list.",
            "Center pane shows chronological Gmail conversation.",
            "Right pane shows customer + ticket context.",
            "Use AI Assist / Escalate / Mark resolved / composer Send as needed.",
            "Deep-link from ticket detail via “Open in Inbox”.",
        ],
    )
    add_heading(doc, "9.4 Confidence Policy (MVP)", 2)
    add_table(
        doc,
        ["Confidence", "Action"],
        [
            [">= 85%", "Auto-reply with grounded answer"],
            ["60% – 84%", "Ask customer for clarification"],
            ["< 60% or critical", "Escalate; send safe escalation message"],
        ],
    )

    # 10. APIs
    add_heading(doc, "10. Main API Endpoints", 1)
    add_table(
        doc,
        ["Method", "Endpoint", "Purpose"],
        [
            ["GET", "/api/gmail/inbox", "List support tickets / connection state"],
            ["GET", "/api/gmail/threads/[threadId]", "Load full Gmail thread"],
            ["POST", "/api/gmail/threads/[threadId]/reply", "Send reply in thread"],
            ["POST", "/api/gmail/threads/[threadId]/process", "Run AI process / escalate / resolve"],
            ["POST", "/api/gmail/poll", "Poll for new activity (live mode)"],
            ["GET", "/api/tickets", "List created tickets"],
            ["GET", "/api/tickets?id=...", "Get one created ticket"],
            ["POST", "/api/tickets", "Create ticket + notify + run agent"],
            ["POST", "/api/tickets/[id]/update", "Send AI/manual customer update"],
            ["GET", "/api/health", "Swytchcode/tooling health"],
        ],
    )

    # 11. Coding highlights
    add_heading(doc, "11. Important Coding Highlights", 1)
    add_bullets(
        doc,
        [
            "src/lib/gmail/create.ts — createSupportTicketFromInput(): notification + agent + auto update.",
            "src/lib/agent/answer.ts — runSupportAgent(): full decision pipeline.",
            "src/lib/gmail/client.ts — uses Swytchcode OAuth in live mode (no fake Bearer override).",
            "src/lib/swytchcode/exec.ts — allowlisted exec with dry-run/live/simulation modes.",
            "src/lib/gmail/store.ts — processed message IDs + created ticket registry.",
            "src/components/tickets/TicketsWorkspace.tsx — create form + clickable ticket rows.",
            "src/app/(product)/tickets/[id]/page.tsx — ticket detail with AI update actions.",
            "src/components/inbox/InboxWorkspace.tsx — 3-pane live inbox UI.",
        ],
    )

    # 12. Demo script
    add_heading(doc, "12. Recommended Demo Script for Judges", 1)
    add_numbered(
        doc,
        [
            "Show swytchcode auth status → Gmail connected.",
            "Open Tickets → create ticket with a real customer email and a Wi-Fi/device issue.",
            "Show success card: Ticket ID, customer email, Notification Sent, AI confidence.",
            "Open customer Gmail and show “[ResolveAI] Ticket created…” and AI follow-up.",
            "Click ticket in ResolveAI → show email + activity + knowledge.",
            "Click “Send AI update to customer” → show new update email arrives.",
            "Open Inbox → same thread appears and conversation is visible.",
            "Mention confidence policy and that unsafe/low-confidence cases escalate.",
        ],
    )

    # 13. Testing
    add_heading(doc, "13. Verification Checklist", 1)
    add_bullets(
        doc,
        [
            "swytchcode list tooling includes required Gmail methods.",
            "swytchcode auth status shows Gmail connected.",
            "POST /api/tickets returns ok=true, notification.sent=true.",
            "Customer receives notification email.",
            "GET /api/gmail/inbox includes created tickets.",
            "Ticket detail page opens on row click.",
            "POST /api/tickets/[id]/update sends AI update email.",
            "No secrets committed (.env.local ignored).",
        ],
    )

    # 14. Limitations
    add_heading(doc, "14. Current Limitations and Future Work", 1)
    add_para(doc, "Limitations:", bold=True)
    add_bullets(
        doc,
        [
            "Knowledge is local markdown, not Notion yet.",
            "Engineering tickets are Gmail/Intercom-oriented; Jira/GitHub not fully wired.",
            "Customer updates currently use Gmail send (Resend not yet integrated).",
            "Polling is used instead of Gmail webhooks for MVP.",
        ],
    )
    add_para(doc, "Future work:", bold=True)
    add_bullets(
        doc,
        [
            "Integrate Notion as live knowledge source via Swytchcode.",
            "Create Jira issues on escalation.",
            "Use Resend for transactional customer updates while keeping Gmail as inbox.",
            "Use GitHub for engineering coordination / issue linking.",
            "Add webhook-based near-real-time inbox updates.",
        ],
    )

    # 15. Conclusion
    add_heading(doc, "15. Conclusion", 1)
    add_para(
        doc,
        "ResolveAI demonstrates an end-to-end AI support workflow on top of Swytchcode and "
        "Gmail: create a ticket from user input, notify the customer by email, analyze the "
        "issue with a grounded agent, automatically send an AI update, and allow operators "
        "to open the ticket later and send additional AI updates—all without exposing Gmail "
        "credentials in the application UI.",
    )

    add_para(doc, "")
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end.add_run("— End of Document —")
    set_run_font(run, size=11, bold=True, color=RGBColor(100, 100, 100))

    out = Path(__file__).resolve().parent.parent / "docs" / "ResolveAI_Project_Submission_Documentation.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(str(out))


if __name__ == "__main__":
    main()
